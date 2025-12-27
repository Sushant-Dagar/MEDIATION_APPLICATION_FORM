"""
Flask Web Application for Mediation Application Form Generator
Allows users to download the generated MS Word document
"""

from flask import Flask, render_template, send_file, request
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

app = Flask(__name__)


def set_cell_borders(cell, border_color="000000", border_size="4"):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_formatted_paragraph(cell, text, bold=False, underline=False, font_size=11, font_name="Times New Roman"):
    """Add a formatted paragraph to a cell."""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(font_size)
    run.font.name = font_name
    return para


def add_new_paragraph(cell, text, bold=False, underline=False, font_size=11, font_name="Times New Roman"):
    """Add a new paragraph to a cell."""
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(font_size)
    run.font.name = font_name
    return para


def create_mediation_form():
    """Create the Mediation Application Form document and return as BytesIO."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # HEADER SECTION
    header1 = doc.add_paragraph()
    header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header1.paragraph_format.space_after = Pt(3)
    header1.paragraph_format.space_before = Pt(0)
    header1.paragraph_format.line_spacing = 1.15
    run1 = header1.add_run("FORM 'A'")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.name = "Times New Roman"

    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.paragraph_format.space_after = Pt(3)
    header2.paragraph_format.space_before = Pt(0)
    header2.paragraph_format.line_spacing = 1.15
    run2 = header2.add_run("MEDIATION APPLICATION FORM")
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.name = "Times New Roman"

    header3 = doc.add_paragraph()
    header3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header3.paragraph_format.space_after = Pt(3)
    header3.paragraph_format.space_before = Pt(0)
    header3.paragraph_format.line_spacing = 1.15
    run3 = header3.add_run("[REFER RULE 3(1)]")
    run3.bold = True
    run3.font.size = Pt(12)
    run3.font.name = "Times New Roman"

    header4 = doc.add_paragraph()
    header4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header4.paragraph_format.space_after = Pt(3)
    header4.paragraph_format.space_before = Pt(0)
    header4.paragraph_format.line_spacing = 1.15
    run4 = header4.add_run("Mumbai District Legal Services Authority")
    run4.font.size = Pt(12)
    run4.font.name = "Times New Roman"

    header5 = doc.add_paragraph()
    header5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header5.paragraph_format.space_after = Pt(12)
    header5.paragraph_format.space_before = Pt(0)
    header5.paragraph_format.line_spacing = 1.15
    run5 = header5.add_run("City Civil Court, Mumbai")
    run5.font.size = Pt(12)
    run5.font.name = "Times New Roman"

    # DETAILS OF PARTIES SECTION
    details_header = doc.add_paragraph()
    details_header.paragraph_format.space_after = Pt(6)
    details_header.paragraph_format.space_before = Pt(6)
    details_header.paragraph_format.line_spacing = 1.15
    run_details = details_header.add_run("DETAILS OF PARTIES:")
    run_details.bold = True
    run_details.font.size = Pt(11)
    run_details.font.name = "Times New Roman"

    # Create main table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(1)
    table.columns[1].width = Cm(4)
    table.columns[2].width = Cm(10)

    # Row 1: Name of Applicant
    row = table.rows[0]
    add_formatted_paragraph(row.cells[0], "1")
    row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_formatted_paragraph(row.cells[1], "Name of", bold=True)
    add_new_paragraph(row.cells[1], "Applicant", bold=True)
    add_formatted_paragraph(row.cells[2], "{{client_name}}")
    for cell in row.cells:
        set_cell_borders(cell)

    # Row 2: Address header (merged)
    row2 = table.add_row()
    row2.cells[0].merge(row2.cells[2])
    add_formatted_paragraph(row2.cells[0], "Address and contact details of Applicant", bold=True, underline=True)
    set_cell_borders(row2.cells[0])

    # Row 3: Address
    row3 = table.add_row()
    add_formatted_paragraph(row3.cells[0], "1")
    row3.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_formatted_paragraph(row3.cells[1], "Address", bold=True)
    cell_addr = row3.cells[2]
    add_formatted_paragraph(cell_addr, "REGISTERED ADDRESS:", bold=True)
    add_new_paragraph(cell_addr, "{{branch_address}}")
    add_new_paragraph(cell_addr, "CORRESPONDENCE BRANCH ADDRESS:", bold=True)
    add_new_paragraph(cell_addr, "{{branch_address}}")
    for cell in row3.cells:
        set_cell_borders(cell)

    # Row 4-6: Contact details
    for label, value in [("Telephone No.", "{{mobile}}"), ("Mobile No.", ""), ("Email ID", "info@kslegal.co.in")]:
        row_n = table.add_row()
        add_formatted_paragraph(row_n.cells[0], "")
        add_formatted_paragraph(row_n.cells[1], label, bold=True)
        add_formatted_paragraph(row_n.cells[2], value)
        if label == "Email ID":
            for run in row_n.cells[2].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
        for cell in row_n.cells:
            set_cell_borders(cell)

    # OPPOSITE PARTY SECTION
    row7 = table.add_row()
    add_formatted_paragraph(row7.cells[0], "2")
    row7.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row7.cells[1].merge(row7.cells[2])
    add_formatted_paragraph(row7.cells[1], "Name, Address and Contact details of Opposite Party:", bold=True)
    for cell in [row7.cells[0], row7.cells[1]]:
        set_cell_borders(cell)

    row8 = table.add_row()
    row8.cells[0].merge(row8.cells[2])
    add_formatted_paragraph(row8.cells[0], "Address and contact details of Defendant/s", bold=True, underline=True)
    set_cell_borders(row8.cells[0])

    # Defendant Name
    row9 = table.add_row()
    add_formatted_paragraph(row9.cells[0], "")
    add_formatted_paragraph(row9.cells[1], "Name", bold=True)
    add_formatted_paragraph(row9.cells[2], "{{customer_name}}")
    for cell in row9.cells:
        set_cell_borders(cell)

    # Defendant Address
    row10 = table.add_row()
    add_formatted_paragraph(row10.cells[0], "")
    add_formatted_paragraph(row10.cells[1], "Address", bold=True)
    cell_def_addr = row10.cells[2]
    add_formatted_paragraph(cell_def_addr, "REGISTERED ADDRESS:", bold=True)
    add_new_paragraph(cell_def_addr, '{% if address1 and address1 != "" %}{{address1}} {% else %} ________________ {%')
    add_new_paragraph(cell_def_addr, "endif %}")
    add_new_paragraph(cell_def_addr, "CORRESPONDENCE ADDRESS:", bold=True)
    add_new_paragraph(cell_def_addr, '{% if address1 and address1 != "" %}{{address1}} {% else %} ________________ {%')
    add_new_paragraph(cell_def_addr, "endif %}")
    for cell in row10.cells:
        set_cell_borders(cell)

    # Defendant contact details
    for label in ["Telephone No.", "Mobile No.", "Email ID"]:
        row_n = table.add_row()
        add_formatted_paragraph(row_n.cells[0], "")
        add_formatted_paragraph(row_n.cells[1], label, bold=True)
        add_formatted_paragraph(row_n.cells[2], "")
        for cell in row_n.cells:
            set_cell_borders(cell)

    # DETAILS OF DISPUTE SECTION
    dispute_header = doc.add_paragraph()
    dispute_header.paragraph_format.space_after = Pt(6)
    dispute_header.paragraph_format.space_before = Pt(12)
    dispute_header.paragraph_format.line_spacing = 1.15
    run_dispute = dispute_header.add_run("DETAILS OF DISPUTE:")
    run_dispute.bold = True
    run_dispute.font.size = Pt(11)
    run_dispute.font.name = "Times New Roman"

    dispute_table = doc.add_table(rows=1, cols=1)
    dispute_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_d1 = dispute_table.rows[0]
    cell_d1 = row_d1.cells[0]
    para_d1 = cell_d1.paragraphs[0]
    para_d1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_d1.paragraph_format.space_before = Pt(4)
    para_d1.paragraph_format.space_after = Pt(4)
    para_d1.paragraph_format.line_spacing = 1.15
    run_d1 = para_d1.add_run("THE COMM. COURTS (PRE-INSTITUTION………SETTLEMENT) RULES,2018")
    run_d1.bold = True
    run_d1.underline = True
    run_d1.font.size = Pt(11)
    run_d1.font.name = "Times New Roman"
    set_cell_borders(cell_d1)

    row_d2 = dispute_table.add_row()
    cell_d2 = row_d2.cells[0]
    add_formatted_paragraph(cell_d2, "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):", bold=True)
    set_cell_borders(cell_d2)

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


@app.route('/')
def index():
    """Render the home page."""
    return render_template('index.html')


@app.route('/download')
def download():
    """Generate and download the Word document."""
    file_stream = create_mediation_form()
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='mediation_application_form.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/preview')
def preview():
    """Show a preview of the form structure."""
    return render_template('preview.html')


if __name__ == '__main__':
    app.run(debug=True)

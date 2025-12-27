"""
Script to create MS Word document replicating the Mediation Application Form (FORM 'A')
from the PDF structure.

Requirements:
    pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE


def set_cell_shading(cell, color="FFFF00"):
    """Set cell background color (yellow by default)."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, border_color="000000", border_size="4"):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_formatted_paragraph(cell, text, bold=False, underline=False, font_size=11, font_name="Times New Roman"):
    """Add a formatted paragraph to a cell."""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.clear()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(font_size)
    run.font.name = font_name
    return para


def add_new_paragraph(cell, text, bold=False, underline=False, font_size=11, font_name="Times New Roman"):
    """Add a new paragraph to a cell."""
    para = cell.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(font_size)
    run.font.name = font_name
    return para


def create_mediation_form():
    """Create the Mediation Application Form document."""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ==================== HEADER SECTION ====================
    # FORM 'A'
    header1 = doc.add_paragraph()
    header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header1.paragraph_format.space_after = Pt(3)
    header1.paragraph_format.space_before = Pt(0)
    header1.paragraph_format.line_spacing = 1.15
    run1 = header1.add_run("FORM 'A'")
    run1.bold = True
    run1.font.size = Pt(12)
    run1.font.name = "Times New Roman"

    # MEDIATION APPLICATION FORM
    header2 = doc.add_paragraph()
    header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header2.paragraph_format.space_after = Pt(3)
    header2.paragraph_format.space_before = Pt(0)
    header2.paragraph_format.line_spacing = 1.15
    run2 = header2.add_run("MEDIATION APPLICATION FORM")
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.name = "Times New Roman"

    # [REFER RULE 3(1)]
    header3 = doc.add_paragraph()
    header3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header3.paragraph_format.space_after = Pt(3)
    header3.paragraph_format.space_before = Pt(0)
    header3.paragraph_format.line_spacing = 1.15
    run3 = header3.add_run("[REFER RULE 3(1)]")
    run3.bold = True
    run3.font.size = Pt(12)
    run3.font.name = "Times New Roman"

    # Mumbai District Legal Services Authority
    header4 = doc.add_paragraph()
    header4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header4.paragraph_format.space_after = Pt(3)
    header4.paragraph_format.space_before = Pt(0)
    header4.paragraph_format.line_spacing = 1.15
    run4 = header4.add_run("Mumbai District Legal Services Authority")
    run4.font.size = Pt(12)
    run4.font.name = "Times New Roman"

    # City Civil Court, Mumbai
    header5 = doc.add_paragraph()
    header5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header5.paragraph_format.space_after = Pt(12)
    header5.paragraph_format.space_before = Pt(0)
    header5.paragraph_format.line_spacing = 1.15
    run5 = header5.add_run("City Civil Court, Mumbai")
    run5.font.size = Pt(12)
    run5.font.name = "Times New Roman"

    # ==================== DETAILS OF PARTIES SECTION ====================
    details_header = doc.add_paragraph()
    details_header.paragraph_format.space_after = Pt(6)
    details_header.paragraph_format.space_before = Pt(6)
    details_header.paragraph_format.line_spacing = 1.15
    run_details = details_header.add_run("DETAILS OF PARTIES:")
    run_details.bold = True
    run_details.font.size = Pt(11)
    run_details.font.name = "Times New Roman"

    # Create main table (3 columns: number, label, value)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    table.columns[0].width = Cm(1)
    table.columns[1].width = Cm(4)
    table.columns[2].width = Cm(10)

    # Row 1: Name of Applicant
    row = table.rows[0]
    cell0 = row.cells[0]
    cell1 = row.cells[1]
    cell2 = row.cells[2]

    add_formatted_paragraph(cell0, "1")
    cell0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    para1 = add_formatted_paragraph(cell1, "Name of", bold=True)
    add_new_paragraph(cell1, "Applicant", bold=True)

    add_formatted_paragraph(cell2, "{{client_name}}")

    # Set borders for first row
    for cell in row.cells:
        set_cell_borders(cell)

    # Row 2: Address and contact details of Applicant (merged)
    row2 = table.add_row()
    row2.cells[0].merge(row2.cells[2])
    merged_cell = row2.cells[0]
    add_formatted_paragraph(merged_cell, "Address and contact details of Applicant", bold=True, underline=True)
    set_cell_borders(merged_cell)

    # Row 3: Address with REGISTERED and CORRESPONDENCE
    row3 = table.add_row()
    add_formatted_paragraph(row3.cells[0], "1")
    row3.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_formatted_paragraph(row3.cells[1], "Address", bold=True)

    # Cell with addresses
    cell_addr = row3.cells[2]
    para_reg = add_formatted_paragraph(cell_addr, "REGISTERED ADDRESS:", bold=True)
    add_new_paragraph(cell_addr, "{{branch_address}}")
    add_new_paragraph(cell_addr, "CORRESPONDENCE BRANCH ADDRESS:", bold=True)
    add_new_paragraph(cell_addr, "{{branch_address}}")

    for cell in row3.cells:
        set_cell_borders(cell)

    # Row 4: Telephone No.
    row4 = table.add_row()
    add_formatted_paragraph(row4.cells[0], "")
    add_formatted_paragraph(row4.cells[1], "Telephone No.", bold=True)
    add_formatted_paragraph(row4.cells[2], "{{mobile}}")
    for cell in row4.cells:
        set_cell_borders(cell)

    # Row 5: Mobile No.
    row5 = table.add_row()
    add_formatted_paragraph(row5.cells[0], "")
    add_formatted_paragraph(row5.cells[1], "Mobile No.", bold=True)
    add_formatted_paragraph(row5.cells[2], "")
    for cell in row5.cells:
        set_cell_borders(cell)

    # Row 6: Email ID
    row6 = table.add_row()
    add_formatted_paragraph(row6.cells[0], "")
    add_formatted_paragraph(row6.cells[1], "Email ID", bold=True)
    para_email = add_formatted_paragraph(row6.cells[2], "info@kslegal.co.in")
    # Make email appear as hyperlink (blue and underlined)
    for run in row6.cells[2].paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.underline = True
    for cell in row6.cells:
        set_cell_borders(cell)

    # ==================== OPPOSITE PARTY SECTION ====================
    # Row 7: Name, Address and Contact details of Opposite Party
    row7 = table.add_row()
    add_formatted_paragraph(row7.cells[0], "2")
    row7.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row7.cells[1].merge(row7.cells[2])
    add_formatted_paragraph(row7.cells[1], "Name, Address and Contact details of Opposite Party:", bold=True)
    for cell in [row7.cells[0], row7.cells[1]]:
        set_cell_borders(cell)

    # Row 8: Address and contact details of Defendant/s
    row8 = table.add_row()
    row8.cells[0].merge(row8.cells[2])
    add_formatted_paragraph(row8.cells[0], "Address and contact details of Defendant/s", bold=True, underline=True)
    set_cell_borders(row8.cells[0])

    # Row 9: Name
    row9 = table.add_row()
    add_formatted_paragraph(row9.cells[0], "")
    add_formatted_paragraph(row9.cells[1], "Name", bold=True)
    add_formatted_paragraph(row9.cells[2], "{{customer_name}}")
    for cell in row9.cells:
        set_cell_borders(cell)

    # Row 10: Address for defendant
    row10 = table.add_row()
    add_formatted_paragraph(row10.cells[0], "")
    add_formatted_paragraph(row10.cells[1], "Address", bold=True)

    # Cell with addresses (with Jinja template syntax)
    cell_def_addr = row10.cells[2]
    add_formatted_paragraph(cell_def_addr, "REGISTERED ADDRESS:", bold=True)
    add_new_paragraph(cell_def_addr, '{% if address1 and address1 != "" %}{{address1}} {% else %} ________________ {%')
    add_new_paragraph(cell_def_addr, "endif %}")
    add_new_paragraph(cell_def_addr, "CORRESPONDENCE ADDRESS:", bold=True)
    add_new_paragraph(cell_def_addr, '{% if address1 and address1 != "" %}{{address1}} {% else %} ________________ {%')
    add_new_paragraph(cell_def_addr, "endif %}")
    for cell in row10.cells:
        set_cell_borders(cell)

    # Row 11: Telephone No. (defendant)
    row11 = table.add_row()
    add_formatted_paragraph(row11.cells[0], "")
    add_formatted_paragraph(row11.cells[1], "Telephone No.", bold=True)
    add_formatted_paragraph(row11.cells[2], "")
    for cell in row11.cells:
        set_cell_borders(cell)

    # Row 12: Mobile No. (defendant)
    row12 = table.add_row()
    add_formatted_paragraph(row12.cells[0], "")
    add_formatted_paragraph(row12.cells[1], "Mobile No.", bold=True)
    add_formatted_paragraph(row12.cells[2], "")
    for cell in row12.cells:
        set_cell_borders(cell)

    # Row 13: Email ID (defendant)
    row13 = table.add_row()
    add_formatted_paragraph(row13.cells[0], "")
    add_formatted_paragraph(row13.cells[1], "Email ID", bold=True)
    add_formatted_paragraph(row13.cells[2], "")
    for cell in row13.cells:
        set_cell_borders(cell)

    # ==================== DETAILS OF DISPUTE SECTION ====================
    dispute_header = doc.add_paragraph()
    dispute_header.paragraph_format.space_after = Pt(6)
    dispute_header.paragraph_format.space_before = Pt(12)
    dispute_header.paragraph_format.line_spacing = 1.15
    run_dispute = dispute_header.add_run("DETAILS OF DISPUTE:")
    run_dispute.bold = True
    run_dispute.font.size = Pt(11)
    run_dispute.font.name = "Times New Roman"

    # Create dispute table
    dispute_table = doc.add_table(rows=1, cols=1)
    dispute_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 1: THE COMM. COURTS header
    row_d1 = dispute_table.rows[0]
    cell_d1 = row_d1.cells[0]
    para_d1 = cell_d1.paragraphs[0]
    para_d1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_d1.paragraph_format.space_before = Pt(4)
    para_d1.paragraph_format.space_after = Pt(4)
    para_d1.paragraph_format.line_spacing = 1.15
    run_d1 = para_d1.add_run("THE COMM. COURTS (PRE-INSTITUTION………SETTLEMENT) RULES,2018")
    run_d1.bold = True
    run_d1.underline = True
    run_d1.font.size = Pt(11)
    run_d1.font.name = "Times New Roman"
    set_cell_borders(cell_d1)

    # Row 2: Nature of disputes
    row_d2 = dispute_table.add_row()
    cell_d2 = row_d2.cells[0]
    para_d2 = add_formatted_paragraph(cell_d2, "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):", bold=True)
    set_cell_borders(cell_d2)

    # Save the document
    import os
    import sys

    # Handle both Windows and Linux paths
    if sys.platform == "win32":
        output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "mediation_application_form.docx")
    else:
        output_path = "/mnt/c/Projects/Web/MediusTechs/mediation_application_form.docx"

    doc.save(output_path)
    print(f"Document created successfully: {output_path}")
    return output_path


if __name__ == "__main__":
    create_mediation_form()
